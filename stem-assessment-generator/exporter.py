"""
Export module for STEM Assessment Generator
Supports multiple formats: JSON, PDF, Word, Plain Text, and LMS formats
"""
import json
import xml.etree.ElementTree as ET
from datetime import datetime
from typing import List, Dict, Any, Optional
from pathlib import Path
import re
import base64
from io import BytesIO
import subprocess
import tempfile
import os
import logging

# Core dependencies
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

# Word document support
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

import markdown2

from models import Question, AssessmentResponse

# Setup logger
logger = logging.getLogger(__name__)

class ExportConfiguration:
    """Configuration options for export formats"""
    
    def __init__(self):
        self.question_order = "as_generated"  # as_generated, by_difficulty, randomized
        self.include_explanations = True
        self.include_difficulty = True
        self.include_metadata = True
        self.teacher_version = True  # Include answers and explanations
        self.font_size = 12
        self.spacing = 1.5
        self.questions_per_page = 5
        self.split_by = "none"  # none, difficulty, concept

class AssessmentExporter:
    """Main exporter class for multiple formats"""
    
    def __init__(self, config: Optional[ExportConfiguration] = None):
        self.config = config or ExportConfiguration()
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom PDF styles"""
        # Question title style
        self.styles.add(ParagraphStyle(
            name='QuestionTitle',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceBefore=12,
            spaceAfter=6,
            leftIndent=0
        ))
        
        # Option style
        self.styles.add(ParagraphStyle(
            name='Option',
            parent=self.styles['Normal'],
            fontSize=11,
            leftIndent=20,
            spaceBefore=3,
            spaceAfter=3
        ))
        
        # Correct answer style
        self.styles.add(ParagraphStyle(
            name='CorrectOption',
            parent=self.styles['Option'],
            textColor=colors.darkgreen,
            bold=True
        ))
    
    # =============================================================================
    # JSON Export Methods
    # =============================================================================
    
    def export_to_json(self, assessment: AssessmentResponse, output_path: str) -> str:
        """Export assessment to JSON format"""
        
        export_data = {
            "version": "1.0",
            "format": "stem_assessment_json",
            "generated_at": datetime.now().isoformat(),
            "source_file": assessment.source_file,
            "processing_time": assessment.processing_time,
            "statistics": assessment.statistics,
            "metadata": assessment.metadata,
            "configuration": {
                "include_explanations": self.config.include_explanations,
                "include_difficulty": self.config.include_difficulty,
                "teacher_version": self.config.teacher_version
            },
            "questions": []
        }
        
        # Process questions
        for i, question in enumerate(assessment.questions, 1):
            question_data = {
                "id": i,
                "question_text": question.question_text,
                "question_type": question.question_type,
                "difficulty": question.difficulty if self.config.include_difficulty else None,
                "options": question.options,
                "correct_answer": question.correct_answer if self.config.teacher_version else None,
                "explanation": question.explanation if self.config.include_explanations else None,
                "distractors": [
                    {
                        "text": d.text,
                        "misconception_type": d.misconception_type,
                        "closeness_score": d.closeness_score,
                        "explanation": d.explanation
                    } for d in question.distractors
                ] if self.config.teacher_version else None,
                "metadata": question.metadata if self.config.include_metadata else None
            }
            export_data["questions"].append(question_data)
        
        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
        
        return output_path
    
    def export_to_qti(self, assessment: AssessmentResponse, output_path: str) -> str:
        """Export to QTI (Question & Test Interoperability) format"""
        
        # Create QTI XML structure
        root = ET.Element("questestinterop")
        root.set("xmlns", "http://www.imsglobal.org/xsd/ims_qtiasiv1p2")
        
        # Assessment metadata
        assessment_elem = ET.SubElement(root, "assessment")
        assessment_elem.set("ident", f"assessment_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
        assessment_elem.set("title", f"Blockchain Assessment - {assessment.source_file}")
        
        # Assessment metadata
        qtimetadata = ET.SubElement(assessment_elem, "qtimetadata")
        qtimetadatafield = ET.SubElement(qtimetadata, "qtimetadatafield")
        ET.SubElement(qtimetadatafield, "fieldlabel").text = "qmd_assessmenttype"
        ET.SubElement(qtimetadatafield, "fieldentry").text = "Examination"
        
        # Sections
        section = ET.SubElement(assessment_elem, "section")
        section.set("ident", "main_section")
        
        # Questions
        for i, question in enumerate(assessment.questions, 1):
            item = ET.SubElement(section, "item")
            item.set("ident", f"question_{i}")
            item.set("title", f"Question {i}")
            
            # Question metadata
            itemmetadata = ET.SubElement(item, "itemmetadata")
            
            # Question text
            presentation = ET.SubElement(item, "presentation")
            material = ET.SubElement(presentation, "material")
            mattext = ET.SubElement(material, "mattext")
            mattext.set("texttype", "text/html")
            mattext.text = f"<p>{question.question_text}</p>"
            
            # Response options
            response_lid = ET.SubElement(presentation, "response_lid")
            response_lid.set("ident", f"response_{i}")
            response_lid.set("rcardinality", "Single")
            
            render_choice = ET.SubElement(response_lid, "render_choice")
            
            for option_key, option_text in question.options.items():
                response_label = ET.SubElement(render_choice, "response_label")
                response_label.set("ident", option_key)
                
                material = ET.SubElement(response_label, "material")
                mattext = ET.SubElement(material, "mattext")
                mattext.set("texttype", "text/plain")
                mattext.text = option_text
            
            # Correct answer (if teacher version)
            if self.config.teacher_version:
                resprocessing = ET.SubElement(item, "resprocessing")
                outcomes = ET.SubElement(resprocessing, "outcomes")
                decvar = ET.SubElement(outcomes, "decvar")
                decvar.set("maxvalue", "100")
                decvar.set("minvalue", "0")
                decvar.set("varname", "SCORE")
                decvar.set("vartype", "Decimal")
                
                # Find correct option key
                correct_key = None
                for key, text in question.options.items():
                    if text == question.correct_answer:
                        correct_key = key
                        break
                
                if correct_key:
                    respcondition = ET.SubElement(resprocessing, "respcondition")
                    respcondition.set("continue", "No")
                    conditionvar = ET.SubElement(respcondition, "conditionvar")
                    varequal = ET.SubElement(conditionvar, "varequal")
                    varequal.set("respident", f"response_{i}")
                    varequal.text = correct_key
                    
                    setvar = ET.SubElement(respcondition, "setvar")
                    setvar.set("action", "Set")
                    setvar.set("varname", "SCORE")
                    setvar.text = "100"
        
        # Write XML to file
        tree = ET.ElementTree(root)
        tree.write(output_path, encoding='utf-8', xml_declaration=True)
        
        return output_path
    
    def export_to_moodle_xml(self, assessment: AssessmentResponse, output_path: str) -> str:
        """Export to Moodle XML format"""
        
        root = ET.Element("quiz")
        
        # Quiz metadata
        question_category = ET.SubElement(root, "question")
        question_category.set("type", "category")
        category = ET.SubElement(question_category, "category")
        text = ET.SubElement(category, "text")
        text.text = f"$course$/Blockchain Assessment/{assessment.source_file}"
        
        # Questions
        for i, question in enumerate(assessment.questions, 1):
            question_elem = ET.SubElement(root, "question")
            question_elem.set("type", "multichoice")
            
            # Question name
            name = ET.SubElement(question_elem, "name")
            text = ET.SubElement(name, "text")
            text.text = f"Question {i}"
            
            # Question text
            questiontext = ET.SubElement(question_elem, "questiontext")
            questiontext.set("format", "html")
            text = ET.SubElement(questiontext, "text")
            text.text = f"<p>{question.question_text}</p>"
            
            # Default grade
            defaultgrade = ET.SubElement(question_elem, "defaultgrade")
            defaultgrade.text = "1"
            
            # Penalty
            penalty = ET.SubElement(question_elem, "penalty")
            penalty.text = "0.1"
            
            # Hidden
            hidden = ET.SubElement(question_elem, "hidden")
            hidden.text = "0"
            
            # Single answer
            single = ET.SubElement(question_elem, "single")
            single.text = "true"
            
            # Shuffle answers
            shuffleanswers = ET.SubElement(question_elem, "shuffleanswers")
            shuffleanswers.text = "true"
            
            # Answer options
            for option_key, option_text in question.options.items():
                answer = ET.SubElement(question_elem, "answer")
                
                # Determine if correct
                is_correct = option_text == question.correct_answer
                answer.set("fraction", "100" if is_correct else "0")
                answer.set("format", "html")
                
                text = ET.SubElement(answer, "text")
                text.text = f"<p>{option_text}</p>"
                
                # Feedback
                feedback = ET.SubElement(answer, "feedback")
                feedback.set("format", "html")
                feedback_text = ET.SubElement(feedback, "text")
                
                if is_correct and question.explanation:
                    feedback_text.text = f"<p>Correct! {question.explanation}</p>"
                elif not is_correct:
                    # Find misconception explanation
                    distractor_explanation = ""
                    for d in question.distractors:
                        if d.text == option_text and d.explanation:
                            distractor_explanation = d.explanation
                            break
                    feedback_text.text = f"<p>Incorrect. {distractor_explanation}</p>"
        
        # Write XML to file
        tree = ET.ElementTree(root)
        tree.write(output_path, encoding='utf-8', xml_declaration=True)
        
        return output_path
    
    # =============================================================================
    # PDF Export Methods
    # =============================================================================
    
    def export_to_pdf(self, assessment: AssessmentResponse, output_path: str) -> str:
        """Export assessment to PDF format"""
        
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        story = []
        
        # Cover page
        story.extend(self._create_pdf_cover_page(assessment))
        story.append(PageBreak())
        
        # Questions
        story.extend(self._create_pdf_questions(assessment))
        
        # Answer key (if teacher version)
        if self.config.teacher_version:
            story.append(PageBreak())
            story.extend(self._create_pdf_answer_key(assessment))
        
        # Build PDF
        doc.build(story)
        
        return output_path

    def export_to_pdf_buffer(self, questions: List[Dict], include_answers: bool = False) -> BytesIO:
        """Export questions to PDF buffer for immediate download - matches Word document structure"""
        buffer = BytesIO()
        
        try:
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch)
            story = []
            styles = getSampleStyleSheet()
            
            # Create custom styles that match Word document
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=18,
                spaceAfter=20,
                alignment=TA_CENTER,
                textColor=colors.black
            )
            
            subtitle_style = ParagraphStyle(
                'SubtitleStyle',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=20,
                alignment=TA_CENTER,
                fontName='Helvetica-Oblique'
            )
            
            question_heading_style = ParagraphStyle(
                'QuestionHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceBefore=15,
                spaceAfter=8,
                textColor=colors.black
            )
            
            question_text_style = ParagraphStyle(
                'QuestionText',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=10,
                fontName='Helvetica'
            )
            
            option_style = ParagraphStyle(
                'OptionStyle',
                parent=styles['Normal'],
                fontSize=11,
                leftIndent=0,
                spaceAfter=4,
                fontName='Helvetica'
            )
            
            correct_option_style = ParagraphStyle(
                'CorrectOptionStyle',
                parent=styles['Normal'],
                fontSize=11,
                leftIndent=0,
                spaceAfter=4,
                fontName='Helvetica-Bold'
            )
            
            explanation_style = ParagraphStyle(
                'ExplanationStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=8,
                fontName='Helvetica-Oblique',
                leftIndent=0
            )
            
            answer_key_style = ParagraphStyle(
                'AnswerKeyStyle',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6
            )
            
            # Add title and subtitle
            story.append(Paragraph("Assessment Questions", title_style))
            version_text = "Teacher Edition (with answers)" if include_answers else "Student Edition"
            story.append(Paragraph(version_text, subtitle_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Add each question
            for i, question in enumerate(questions):
                # Question heading
                story.append(Paragraph(f"Question {i+1}", question_heading_style))
                
                # Question text
                question_text = question.get('question_text', '')
                story.append(Paragraph(question_text, question_text_style))
                story.append(Spacer(1, 0.1*inch))
                
                # Handle different option formats
                options = question.get('options', {})
                if isinstance(options, dict):
                    # New format with detailed options
                    for key in ['A', 'B', 'C', 'D']:
                        if key in options:
                            option = options[key]
                            if isinstance(option, dict):
                                option_text = f"{key}. {option.get('text', '')}"
                                is_correct = option.get('is_correct', False)
                            else:
                                option_text = f"{key}. {option}"
                                is_correct = (option == question.get('correct_answer', ''))
                            
                            if include_answers and is_correct:
                                # Add checkmark and "Correct Answer" text for teacher version
                                enhanced_text = f"<b>{option_text} ✓ (Correct Answer)</b>"
                                story.append(Paragraph(enhanced_text, correct_option_style))
                            else:
                                story.append(Paragraph(option_text, option_style))
                else:
                    # Legacy format
                    correct_answer = question.get('correct_answer', '')
                    distractors = question.get('distractors', [])
                    
                    # Add correct answer as option A
                    option_text = f"A. {correct_answer}"
                    if include_answers:
                        enhanced_text = f"<b>{option_text} ✓ (Correct Answer)</b>"
                        story.append(Paragraph(enhanced_text, correct_option_style))
                    else:
                        story.append(Paragraph(option_text, option_style))
                    
                    # Add distractors
                    for j, distractor in enumerate(distractors[:3]):
                        letter = chr(66 + j)  # B, C, D
                        d_text = distractor if isinstance(distractor, str) else distractor.get('text', '')
                        option_text = f"{letter}. {d_text}"
                        story.append(Paragraph(option_text, option_style))
                
                # Add explanation for teacher version
                if include_answers:
                    story.append(Spacer(1, 0.05*inch))
                    explanation_text = question.get('explanation', 'See detailed explanations in the teaching materials.')
                    story.append(Paragraph(f"<b>Explanation:</b> <i>{explanation_text}</i>", explanation_style))
                
                # Add space between questions
                story.append(Spacer(1, 0.15*inch))
                story.append(Spacer(1, 0.15*inch))
                
                # Page break after every 3 questions (except last)
                if (i + 1) % 3 == 0 and i < len(questions) - 1:
                    story.append(PageBreak())
            
            # Add answer key section for teacher version
            if include_answers:
                story.append(PageBreak())
                story.append(Paragraph("Answer Key", title_style))
                story.append(Spacer(1, 0.2*inch))
                
                for i, question in enumerate(questions):
                    # Find correct answer
                    correct_text = ""
                    options = question.get('options', {})
                    
                    if isinstance(options, dict):
                        for key, option in options.items():
                            if isinstance(option, dict) and option.get('is_correct'):
                                correct_text = f"{key}. {option.get('text', '')}"
                                break
                    else:
                        correct_text = f"A. {question.get('correct_answer', '')}"
                    
                    # Add answer entry
                    answer_text = f"<b>Question {i+1}:</b> {correct_text}"
                    
                    # Add difficulty if available
                    difficulty = question.get('difficulty', '')
                    if difficulty:
                        answer_text += f" <i>(Difficulty: {difficulty})</i>"
                    
                    story.append(Paragraph(answer_text, answer_key_style))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            # If PDF generation fails, create a simple error PDF
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = [
                Paragraph("PDF Generation Error", styles['Title']),
                Spacer(1, 0.5*inch),
                Paragraph(f"Error: {str(e)}", styles['Normal']),
                Spacer(1, 0.2*inch),
                Paragraph("Please try again or contact support.", styles['Normal'])
            ]
            doc.build(story)
            buffer.seek(0)
            return buffer

    def export_to_docx_buffer(self, questions: List[Dict], include_answers: bool = False) -> BytesIO:
        """Export questions to Word document buffer for immediate download"""
        buffer = BytesIO()
        
        try:
            doc = Document()
            
            # Set document styles
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            
            # Title
            title_paragraph = doc.add_heading('Assessment Questions', 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Subtitle
            version_text = "Teacher Edition (with answers)" if include_answers else "Student Edition"
            subtitle = doc.add_paragraph(version_text)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].italic = True
            
            doc.add_paragraph()  # Add space
            
            # Add each question
            for i, question in enumerate(questions):
                # Question heading
                question_heading = doc.add_heading(f'Question {i+1}', level=2)
                
                # Question text
                question_text = question.get('question_text', '')
                doc.add_paragraph(question_text)
                
                # Add some space
                doc.add_paragraph()
                
                # Handle different option formats
                options = question.get('options', {})
                if isinstance(options, dict):
                    # New format with detailed options
                    for key in ['A', 'B', 'C', 'D']:
                        if key in options:
                            option = options[key]
                            if isinstance(option, dict):
                                option_text = f"{key}. {option.get('text', '')}"
                                is_correct = option.get('is_correct', False)
                            else:
                                option_text = f"{key}. {option}"
                                is_correct = (option == question.get('correct_answer', ''))
                            
                            option_para = doc.add_paragraph(option_text)
                            
                            # Mark correct answer if teacher version
                            if include_answers and is_correct:
                                # Make the option bold and add checkmark
                                option_para.runs[0].bold = True
                                option_para.add_run(' ✓').bold = True
                                option_para.add_run(' (Correct Answer)').italic = True
                else:
                    # Legacy format
                    correct_answer = question.get('correct_answer', '')
                    distractors = question.get('distractors', [])
                    
                    # Add correct answer as option A
                    option_text = f"A. {correct_answer}"
                    option_para = doc.add_paragraph(option_text)
                    if include_answers:
                        option_para.runs[0].bold = True
                        option_para.add_run(' ✓ (Correct Answer)').italic = True
                    
                    # Add distractors
                    for j, distractor in enumerate(distractors[:3]):
                        letter = chr(66 + j)  # B, C, D
                        d_text = distractor if isinstance(distractor, str) else distractor.get('text', '')
                        option_text = f"{letter}. {d_text}"
                        doc.add_paragraph(option_text)
                
                # Add explanation if teacher version
                if include_answers:
                    doc.add_paragraph()
                    explanation_para = doc.add_paragraph()
                    explanation_para.add_run('Explanation: ').bold = True
                    
                    explanation_text = question.get('explanation', 'See detailed explanations in the teaching materials.')
                    explanation_para.add_run(explanation_text).italic = True
                
                # Add space between questions
                doc.add_paragraph()
                doc.add_paragraph()
                
                # Page break after every 3 questions (except last)
                if (i + 1) % 3 == 0 and i < len(questions) - 1:
                    doc.add_page_break()
            
            # Add answer key section for teacher version
            if include_answers:
                doc.add_page_break()
                doc.add_heading('Answer Key', level=1)
                doc.add_paragraph()
                
                for i, question in enumerate(questions):
                    # Find correct answer
                    correct_text = ""
                    options = question.get('options', {})
                    
                    if isinstance(options, dict):
                        for key, option in options.items():
                            if isinstance(option, dict) and option.get('is_correct'):
                                correct_text = f"{key}. {option.get('text', '')}"
                                break
                    else:
                        correct_text = f"A. {question.get('correct_answer', '')}"
                    
                    # Add answer entry
                    answer_para = doc.add_paragraph()
                    answer_para.add_run(f"Question {i+1}: ").bold = True
                    answer_para.add_run(correct_text)
                    
                    # Add difficulty if available
                    difficulty = question.get('difficulty', '')
                    if difficulty:
                        answer_para.add_run(f" (Difficulty: {difficulty})").italic = True
            
            # Save to buffer
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            # If Word generation fails, create a simple error document
            buffer = BytesIO()
            doc = Document()
            doc.add_heading('Document Generation Error', 0)
            doc.add_paragraph(f"Error: {str(e)}")
            doc.add_paragraph("Please try again or contact support.")
            doc.save(buffer)
            buffer.seek(0)
            return buffer
    
    def export_word_to_pdf_buffer(self, assessment: AssessmentResponse, 
                                 teacher_version: bool = True, 
                                 include_explanations: bool = True) -> BytesIO:
        """Convert Word document to PDF using LibreOffice CLI"""
        try:
            # First create Word document in temp file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                word_buffer = self.export_to_docx_buffer(
                    [q.dict() for q in assessment.questions], 
                    include_answers=teacher_version
                )
                temp_docx.write(word_buffer.read())
                temp_docx_path = temp_docx.name
            
            # Create temp directory for PDF output
            temp_dir = tempfile.mkdtemp()
            
            try:
                # Use LibreOffice to convert Word to PDF
                cmd = [
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', temp_dir, temp_docx_path
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                
                if result.returncode != 0:
                    logger.warning(f"LibreOffice conversion failed: {result.stderr}")
                    # Fallback to ReportLab
                    return self.export_assessment_to_pdf_buffer(assessment, teacher_version, include_explanations)
                
                # Find the generated PDF file
                pdf_filename = os.path.splitext(os.path.basename(temp_docx_path))[0] + '.pdf'
                pdf_path = os.path.join(temp_dir, pdf_filename)
                
                if os.path.exists(pdf_path):
                    # Read PDF into buffer
                    with open(pdf_path, 'rb') as pdf_file:
                        pdf_buffer = BytesIO(pdf_file.read())
                    return pdf_buffer
                else:
                    logger.warning("PDF file not found after conversion")
                    # Fallback to ReportLab
                    return self.export_assessment_to_pdf_buffer(assessment, teacher_version, include_explanations)
                    
            finally:
                # Clean up temp files
                try:
                    os.unlink(temp_docx_path)
                    if os.path.exists(temp_dir):
                        for file in os.listdir(temp_dir):
                            os.unlink(os.path.join(temp_dir, file))
                        os.rmdir(temp_dir)
                except:
                    pass
                    
        except Exception as e:
            logger.error(f"Word to PDF conversion failed: {e}")
            # Fallback to ReportLab method
            return self.export_assessment_to_pdf_buffer(assessment, teacher_version, include_explanations)
    
    def export_assessment_to_pdf_buffer(self, assessment: AssessmentResponse, 
                                      teacher_version: bool = True, 
                                      include_explanations: bool = True) -> BytesIO:
        """Export AssessmentResponse to PDF buffer for immediate download"""
        buffer = BytesIO()
        
        try:
            # Create document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Create styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=20,
                alignment=TA_CENTER
            )
            
            question_style = ParagraphStyle(
                'QuestionStyle',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
                fontName='Helvetica-Bold'
            )
            
            option_style = ParagraphStyle(
                'OptionStyle',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6,
                leftIndent=20
            )
            
            answer_style = ParagraphStyle(
                'AnswerStyle',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                leftIndent=20,
                fontName='Helvetica-Bold',
                textColor=colors.darkgreen
            )
            
            explanation_style = ParagraphStyle(
                'ExplanationStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=15,
                leftIndent=20,
                fontName='Helvetica-Oblique',
                textColor=colors.darkblue
            )
            
            # Build content
            content = []
            
            # Title
            version_text = "Teacher Version" if teacher_version else "Student Version"
            title = f"STEM Assessment - {version_text}"
            content.append(Paragraph(title, title_style))
            content.append(Spacer(1, 20))
            
            # Assessment info
            if hasattr(assessment, 'source_file') and assessment.source_file:
                content.append(Paragraph(f"<b>Source:</b> {assessment.source_file}", styles['Normal']))
            content.append(Paragraph(f"<b>Generated:</b> {assessment.generated_at.strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
            content.append(Paragraph(f"<b>Total Questions:</b> {len(assessment.questions)}", styles['Normal']))
            content.append(Spacer(1, 20))
            
            # Questions
            for i, question in enumerate(assessment.questions, 1):
                # Question number and text
                question_text = f"<b>Question {i}:</b> {question.question_text}"
                content.append(Paragraph(question_text, question_style))
                
                # Options (use the new options property)
                try:
                    options = question.options  # This should work now with our property
                    for key in ['A', 'B', 'C', 'D']:
                        if key in options:
                            option_text = f"{key}. {options[key]}"
                            
                            # Highlight correct answer for teacher version
                            if teacher_version and key == 'A':  # Correct answer is always A in our structure
                                content.append(Paragraph(f"<b>{option_text} ✓ (Correct)</b>", answer_style))
                            else:
                                content.append(Paragraph(option_text, option_style))
                except Exception as e:
                    # Fallback to old format
                    content.append(Paragraph(f"A. {question.correct_answer}", answer_style if teacher_version else option_style))
                    for j, distractor in enumerate(question.distractors[:3]):
                        letter = ['B', 'C', 'D'][j]
                        content.append(Paragraph(f"{letter}. {distractor.text}", option_style))
                
                # Add explanations for teacher version
                if teacher_version and include_explanations and question.explanation:
                    content.append(Paragraph(f"<b>Explanation:</b> {question.explanation}", explanation_style))
                
                # Add misconception details if available
                if teacher_version and include_explanations:
                    try:
                        for j, distractor in enumerate(question.distractors):
                            if distractor.misconception_type and distractor.explanation:
                                letter = ['B', 'C', 'D'][j]
                                misconception_text = f"<b>Option {letter} Misconception:</b> {distractor.misconception_type} - {distractor.explanation}"
                                content.append(Paragraph(misconception_text, explanation_style))
                    except:
                        pass
                
                content.append(Spacer(1, 15))
            
            # Build PDF
            doc.build(content)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            # Create a simple error PDF
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            content = [Paragraph(f"PDF Generation Error: {str(e)}", getSampleStyleSheet()['Normal'])]
            doc.build(content)
            buffer.seek(0)
            return buffer
    
    def export_to_docx_buffer_assessment(self, assessment: AssessmentResponse, 
                                       include_answers: bool = False,
                                       include_explanations: bool = True) -> BytesIO:
        """Export AssessmentResponse to Word document buffer with detailed explanations"""
        buffer = BytesIO()
        
        try:
            doc = Document()
            
            # Set document styles
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            
            # Title
            version_text = "Teacher Edition (with answers)" if include_answers else "Student Edition"
            title = doc.add_heading(f'Assessment Questions - {version_text}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Assessment metadata
            if hasattr(assessment, 'source_file') and assessment.source_file:
                doc.add_paragraph(f"Source: {assessment.source_file}")
            doc.add_paragraph(f"Generated: {assessment.generated_at.strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph(f"Total Questions: {len(assessment.questions)}")
            doc.add_paragraph()  # Add space
            
            # Process each question
            for i, question in enumerate(assessment.questions, 1):
                # Question heading
                question_heading = doc.add_heading(f'Question {i}', level=2)
                
                # Question text
                doc.add_paragraph(question.question_text)
                doc.add_paragraph()  # Space before options
                
                # Options using the new options property
                try:
                    options = question.options  # This should work with our property
                    for key in ['A', 'B', 'C', 'D']:
                        if key in options:
                            option_text = f"{key}. {options[key]}"
                            option_para = doc.add_paragraph(option_text)
                            
                            # Mark correct answer if teacher version
                            if include_answers and key == 'A':  # Correct answer is always A in our structure
                                option_para.runs[0].bold = True
                                option_para.add_run(' ✓ (Correct Answer)').bold = True
                                
                except Exception as e:
                    # Fallback to old format
                    option_para = doc.add_paragraph(f"A. {question.correct_answer}")
                    if include_answers:
                        option_para.runs[0].bold = True
                        option_para.add_run(' ✓ (Correct Answer)').bold = True
                    
                    for j, distractor in enumerate(question.distractors[:3]):
                        letter = ['B', 'C', 'D'][j]
                        doc.add_paragraph(f"{letter}. {distractor.text}")
                
                doc.add_paragraph()  # Space after options
                
                # Add explanations for teacher version
                if include_answers and include_explanations:
                    if question.explanation:
                        explanation_heading = doc.add_paragraph()
                        explanation_heading.add_run("Explanation:").bold = True
                        doc.add_paragraph(question.explanation)
                    
                    # Add detailed misconception explanations if available
                    try:
                        for j, distractor in enumerate(question.distractors):
                            if distractor.misconception_type and distractor.explanation:
                                letter = ['B', 'C', 'D'][j]
                                misconception_para = doc.add_paragraph()
                                misconception_para.add_run(f"Option {letter} Misconception:").bold = True
                                misconception_para.add_run(f" {distractor.misconception_type} - {distractor.explanation}")
                    except:
                        pass
                
                # Add difficulty if available
                if hasattr(question, 'difficulty') and question.difficulty:
                    difficulty_para = doc.add_paragraph()
                    difficulty_para.add_run(f"Difficulty: {question.difficulty}").italic = True
                
                # Page break after each question (except the last one)
                if i < len(assessment.questions):
                    doc.add_page_break()
            
            # Save to buffer
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logger.error(f"Word document generation failed: {e}")
            # Create simple error document
            buffer = BytesIO()
            doc = Document()
            doc.add_heading('Document Generation Error', 0)
            doc.add_paragraph(f"Error: {str(e)}")
            doc.add_paragraph("Please try again or contact support.")
            doc.save(buffer)
            buffer.seek(0)
            return buffer
    
    def convert_word_to_pdf(self, word_buffer: BytesIO) -> BytesIO:
        """Convert Word document buffer to PDF using LibreOffice"""
        import tempfile
        import subprocess
        import os
        
        try:
            # Create temporary files
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                temp_docx.write(word_buffer.read())
                temp_docx_path = temp_docx.name
            
            temp_dir = tempfile.mkdtemp()
            
            # Convert using LibreOffice
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', temp_dir, temp_docx_path
            ], check=True, capture_output=True)
            
            # Read the generated PDF
            pdf_filename = os.path.splitext(os.path.basename(temp_docx_path))[0] + '.pdf'
            pdf_path = os.path.join(temp_dir, pdf_filename)
            
            pdf_buffer = BytesIO()
            with open(pdf_path, 'rb') as pdf_file:
                pdf_buffer.write(pdf_file.read())
            
            # Cleanup
            os.unlink(temp_docx_path)
            os.unlink(pdf_path)
            os.rmdir(temp_dir)
            
            pdf_buffer.seek(0)
            return pdf_buffer
            
        except subprocess.CalledProcessError:
            logger.warning("LibreOffice conversion failed, falling back to ReportLab")
            # Fallback to our existing PDF generation
            word_buffer.seek(0)  # Reset buffer
            return self._fallback_pdf_generation(word_buffer)
        except Exception as e:
            logger.error(f"Word to PDF conversion failed: {e}")
            raise Exception(f"PDF conversion failed: {str(e)}")
    
    def _fallback_pdf_generation(self, word_buffer: BytesIO) -> BytesIO:
        """Fallback PDF generation when LibreOffice is not available"""
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        content = [
            Paragraph("Assessment Questions", styles['Title']),
            Spacer(1, 20),
            Paragraph("PDF generated using fallback method.", styles['Normal']),
            Paragraph("For full formatting, please install LibreOffice.", styles['Normal'])
        ]
        
        doc.build(content)
        buffer.seek(0)
        return buffer
    
    def _create_pdf_cover_page(self, assessment: AssessmentResponse) -> List:
        """Create PDF cover page"""
        story = []
        
        # Title
        title = Paragraph("Blockchain Assessment", self.styles['Title'])
        story.append(title)
        story.append(Spacer(1, 0.5*inch))
        
        # Metadata table
        metadata = [
            ['Source File:', assessment.source_file or 'Unknown'],
            ['Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
            ['Total Questions:', str(len(assessment.questions))],
            ['Processing Time:', f"{assessment.processing_time:.2f} seconds"],
            ['Version:', 'Teacher' if self.config.teacher_version else 'Student']
        ]
        
        table = Table(metadata, colWidths=[2*inch, 4*inch])
        table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        story.append(table)
        story.append(Spacer(1, 0.5*inch))
        
        # Statistics
        if assessment.statistics:
            story.append(Paragraph("Generation Statistics", self.styles['Heading2']))
            story.append(Spacer(1, 0.2*inch))
            
            stats_data = []
            for key, value in assessment.statistics.items():
                if isinstance(value, (int, float, str)):
                    stats_data.append([key.replace('_', ' ').title() + ':', str(value)])
            
            if stats_data:
                stats_table = Table(stats_data, colWidths=[2.5*inch, 3.5*inch])
                stats_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ]))
                story.append(stats_table)
        
        return story
    
    def _create_pdf_questions(self, assessment: AssessmentResponse) -> List:
        """Create PDF questions section"""
        story = []
        
        story.append(Paragraph("Questions", self.styles['Heading1']))
        story.append(Spacer(1, 0.3*inch))
        
        for i, question in enumerate(assessment.questions, 1):
            # Question number and difficulty
            difficulty_badge = f" [{question.difficulty}]" if self.config.include_difficulty else ""
            question_title = f"Question {i}{difficulty_badge}"
            story.append(Paragraph(question_title, self.styles['QuestionTitle']))
            
            # Question text
            story.append(Paragraph(question.question_text, self.styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
            
            # Options
            for option_key, option_text in question.options.items():
                is_correct = option_text == question.correct_answer
                style = self.styles['CorrectOption'] if (is_correct and self.config.teacher_version) else self.styles['Option']
                
                option_para = Paragraph(f"{option_key}. {option_text}", style)
                story.append(option_para)
            
            # Explanation (if teacher version)
            if self.config.teacher_version and self.config.include_explanations and question.explanation:
                story.append(Spacer(1, 0.1*inch))
                story.append(Paragraph(f"<b>Explanation:</b> {question.explanation}", self.styles['Normal']))
            
            story.append(Spacer(1, 0.3*inch))
            
            # Page break after certain number of questions
            if i % self.config.questions_per_page == 0 and i < len(assessment.questions):
                story.append(PageBreak())
        
        return story
    
    def _create_pdf_answer_key(self, assessment: AssessmentResponse) -> List:
        """Create PDF answer key section"""
        story = []
        
        story.append(Paragraph("Answer Key", self.styles['Heading1']))
        story.append(Spacer(1, 0.3*inch))
        
        # Create answer table
        answer_data = [['Question', 'Correct Answer', 'Explanation']]
        
        for i, question in enumerate(assessment.questions, 1):
            # Find correct option key
            correct_key = None
            for key, text in question.options.items():
                if text == question.correct_answer:
                    correct_key = key
                    break
            
            answer_row = [
                str(i),
                f"{correct_key}. {question.correct_answer}" if correct_key else question.correct_answer,
                question.explanation or "No explanation provided"
            ]
            answer_data.append(answer_row)
        
        answer_table = Table(answer_data, colWidths=[0.8*inch, 2.5*inch, 3.2*inch])
        answer_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        
        story.append(answer_table)
        
        return story
    
    # =============================================================================
    # Word Document Export
    # =============================================================================
    
    def export_to_word(self, assessment: AssessmentResponse, output_path: str) -> str:
        """Export assessment to Word document"""
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Blockchain Assessment', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_heading('Assessment Information', level=1)
        
        metadata_table = doc.add_table(rows=5, cols=2)
        metadata_table.style = 'Table Grid'
        
        metadata_data = [
            ('Source File:', assessment.source_file or 'Unknown'),
            ('Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('Total Questions:', str(len(assessment.questions))),
            ('Processing Time:', f"{assessment.processing_time:.2f} seconds"),
            ('Version:', 'Teacher' if self.config.teacher_version else 'Student')
        ]
        
        for i, (label, value) in enumerate(metadata_data):
            row = metadata_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            # Make labels bold
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_page_break()
        
        # Questions
        doc.add_heading('Questions', level=1)
        
        for i, question in enumerate(assessment.questions, 1):
            # Question header
            difficulty_badge = f" [{question.difficulty}]" if self.config.include_difficulty else ""
            question_heading = doc.add_heading(f'Question {i}{difficulty_badge}', level=2)
            
            # Question text
            question_para = doc.add_paragraph(question.question_text)
            question_para.style = 'Normal'
            
            # Options table
            options_table = doc.add_table(rows=len(question.options), cols=2)
            options_table.style = 'Table Grid'
            
            for j, (option_key, option_text) in enumerate(question.options.items()):
                row = options_table.rows[j]
                row.cells[0].text = f"{option_key}."
                row.cells[1].text = option_text
                
                # Highlight correct answer (teacher version)
                if self.config.teacher_version and option_text == question.correct_answer:
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
                    row.cells[1].paragraphs[0].runs[0].font.bold = True
                    row.cells[0].paragraphs[0].runs[0].font.color.rgb = colors.darkgreen
                    row.cells[1].paragraphs[0].runs[0].font.color.rgb = colors.darkgreen
            
            # Explanation (teacher version)
            if self.config.teacher_version and self.config.include_explanations and question.explanation:
                explanation_para = doc.add_paragraph()
                explanation_para.add_run('Explanation: ').bold = True
                explanation_para.add_run(question.explanation)
            
            doc.add_paragraph()  # Add spacing
        
        # Answer key (teacher version)
        if self.config.teacher_version:
            doc.add_page_break()
            doc.add_heading('Answer Key', level=1)
            
            answer_table = doc.add_table(rows=len(assessment.questions) + 1, cols=3)
            answer_table.style = 'Table Grid'
            
            # Header row
            header_cells = answer_table.rows[0].cells
            header_cells[0].text = 'Question'
            header_cells[1].text = 'Correct Answer'
            header_cells[2].text = 'Explanation'
            
            for cell in header_cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Answer rows
            for i, question in enumerate(assessment.questions, 1):
                row = answer_table.rows[i]
                
                # Find correct option key
                correct_key = None
                for key, text in question.options.items():
                    if text == question.correct_answer:
                        correct_key = key
                        break
                
                row.cells[0].text = str(i)
                row.cells[1].text = f"{correct_key}. {question.correct_answer}" if correct_key else question.correct_answer
                row.cells[2].text = question.explanation or "No explanation provided"
        
        # Save document
        doc.save(output_path)
        
        return output_path
    
    # =============================================================================
    # Plain Text and Markdown Export
    # =============================================================================
    
    def export_to_text(self, assessment: AssessmentResponse, output_path: str) -> str:
        """Export assessment to plain text format"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            # Header
            f.write("=" * 60 + "\n")
            f.write("BLOCKCHAIN ASSESSMENT\n")
            f.write("=" * 60 + "\n\n")
            
            # Metadata
            f.write(f"Source File: {assessment.source_file or 'Unknown'}\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Total Questions: {len(assessment.questions)}\n")
            f.write(f"Processing Time: {assessment.processing_time:.2f} seconds\n")
            f.write(f"Version: {'Teacher' if self.config.teacher_version else 'Student'}\n")
            f.write("\n" + "-" * 60 + "\n\n")
            
            # Questions
            for i, question in enumerate(assessment.questions, 1):
                difficulty_badge = f" [{question.difficulty}]" if self.config.include_difficulty else ""
                f.write(f"QUESTION {i}{difficulty_badge}\n")
                f.write("-" * 20 + "\n")
                f.write(f"{question.question_text}\n\n")
                
                # Options
                for option_key, option_text in question.options.items():
                    marker = ">>> " if (self.config.teacher_version and option_text == question.correct_answer) else "    "
                    f.write(f"{marker}{option_key}. {option_text}\n")
                
                # Explanation
                if self.config.teacher_version and self.config.include_explanations and question.explanation:
                    f.write(f"\nExplanation: {question.explanation}\n")
                
                f.write("\n" + "~" * 40 + "\n\n")
            
            # Answer key
            if self.config.teacher_version:
                f.write("\n" + "=" * 60 + "\n")
                f.write("ANSWER KEY\n")
                f.write("=" * 60 + "\n\n")
                
                for i, question in enumerate(assessment.questions, 1):
                    # Find correct option key
                    correct_key = None
                    for key, text in question.options.items():
                        if text == question.correct_answer:
                            correct_key = key
                            break
                    
                    f.write(f"{i}. {correct_key} - {question.correct_answer}\n")
                    if question.explanation:
                        f.write(f"   Explanation: {question.explanation}\n")
                    f.write("\n")
        
        return output_path
    
    def export_to_markdown(self, assessment: AssessmentResponse, output_path: str) -> str:
        """Export assessment to Markdown format"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            # Header
            f.write("# Blockchain Assessment\n\n")
            
            # Metadata
            f.write("## Assessment Information\n\n")
            f.write(f"- **Source File:** {assessment.source_file or 'Unknown'}\n")
            f.write(f"- **Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"- **Total Questions:** {len(assessment.questions)}\n")
            f.write(f"- **Processing Time:** {assessment.processing_time:.2f} seconds\n")
            f.write(f"- **Version:** {'Teacher' if self.config.teacher_version else 'Student'}\n\n")
            
            f.write("---\n\n")
            
            # Questions
            f.write("## Questions\n\n")
            
            for i, question in enumerate(assessment.questions, 1):
                difficulty_badge = f" `[{question.difficulty}]`" if self.config.include_difficulty else ""
                f.write(f"### Question {i}{difficulty_badge}\n\n")
                f.write(f"{question.question_text}\n\n")
                
                # Options
                for option_key, option_text in question.options.items():
                    if self.config.teacher_version and option_text == question.correct_answer:
                        f.write(f"- **{option_key}.** **{option_text}** ✅\n")
                    else:
                        f.write(f"- {option_key}. {option_text}\n")
                
                # Explanation
                if self.config.teacher_version and self.config.include_explanations and question.explanation:
                    f.write(f"\n> **Explanation:** {question.explanation}\n")
                
                f.write("\n---\n\n")
            
            # Answer key
            if self.config.teacher_version:
                f.write("## Answer Key\n\n")
                f.write("| Question | Answer | Explanation |\n")
                f.write("|----------|--------|-------------|\n")
                
                for i, question in enumerate(assessment.questions, 1):
                    # Find correct option key
                    correct_key = None
                    for key, text in question.options.items():
                        if text == question.correct_answer:
                            correct_key = key
                            break
                    
                    answer = f"{correct_key}. {question.correct_answer}" if correct_key else question.correct_answer
                    explanation = question.explanation or "No explanation provided"
                    
                    # Escape pipe characters for table
                    answer = answer.replace('|', '\\|')
                    explanation = explanation.replace('|', '\\|')
                    
                    f.write(f"| {i} | {answer} | {explanation} |\n")
        
        return output_path
    
    # =============================================================================
    # LMS-Specific Formats
    # =============================================================================
    
    def export_to_canvas(self, assessment: AssessmentResponse, output_path: str) -> str:
        """Export to Canvas-compatible CSV format"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            # CSV header for Canvas
            f.write("Title,Question Text,Question Type,Points Possible,Correct Comments,Incorrect Comments,Neutral Comments,")
            f.write("Answer 1,Answer 1 Comment,Answer 1 Weight,Answer 2,Answer 2 Comment,Answer 2 Weight,")
            f.write("Answer 3,Answer 3 Comment,Answer 3 Weight,Answer 4,Answer 4 Comment,Answer 4 Weight\n")
            
            for i, question in enumerate(assessment.questions, 1):
                title = f"Question {i}"
                question_text = question.question_text.replace('"', '""')  # Escape quotes
                question_type = "multiple_choice_question"
                points = "1"
                
                # Comments
                correct_comment = question.explanation.replace('"', '""') if question.explanation else ""
                incorrect_comment = "Please review the material and try again."
                neutral_comment = ""
                
                # Answers (ensure we have exactly 4 options)
                options_list = list(question.options.items())
                while len(options_list) < 4:
                    options_list.append(("", ""))
                
                answers = []
                for j, (option_key, option_text) in enumerate(options_list[:4]):
                    is_correct = option_text == question.correct_answer
                    weight = "100" if is_correct else "0"
                    comment = correct_comment if is_correct else incorrect_comment
                    
                    answers.extend([
                        option_text.replace('"', '""'),
                        comment,
                        weight
                    ])
                
                # Write row
                row = [title, question_text, question_type, points, correct_comment, incorrect_comment, neutral_comment] + answers
                csv_row = ','.join(f'"{field}"' for field in row)
                f.write(csv_row + "\n")
        
        return output_path
    
    def export_to_google_forms(self, assessment: AssessmentResponse, output_path: str) -> str:
        """Export to Google Forms-compatible format (CSV)"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("Question,Option A,Option B,Option C,Option D,Correct Answer,Explanation\n")
            
            for question in assessment.questions:
                # Get options in order
                options = ['', '', '', '']  # Default 4 empty options
                correct_position = 0
                
                for i, (key, text) in enumerate(question.options.items()):
                    if i < 4:
                        options[i] = text.replace('"', '""')
                        if text == question.correct_answer:
                            correct_position = i + 1
                
                question_text = question.question_text.replace('"', '""')
                explanation = question.explanation.replace('"', '""') if question.explanation else ""
                
                row = [question_text] + options + [str(correct_position), explanation]
                csv_row = ','.join(f'"{field}"' for field in row)
                f.write(csv_row + "\n")
        
        return output_path

# =============================================================================
# Export Manager Class
# =============================================================================

class ExportManager:
    """Manager class for handling all export operations"""
    
    def __init__(self):
        self.exporters = {}
    
    def export_assessment(self, assessment: AssessmentResponse, format_type: str, 
                         output_path: str, config: Optional[ExportConfiguration] = None) -> str:
        """
        Export assessment to specified format
        
        Args:
            assessment: AssessmentResponse object
            format_type: Export format (json, pdf, word, text, markdown, qti, moodle, canvas, google_forms)
            output_path: Output file path
            config: Export configuration options
        
        Returns:
            Path to exported file
        """
        
        if config is None:
            config = ExportConfiguration()
        
        exporter = AssessmentExporter(config)
        
        format_methods = {
            'json': exporter.export_to_json,
            'pdf': exporter.export_to_pdf,
            'word': exporter.export_to_word,
            'text': exporter.export_to_text,
            'markdown': exporter.export_to_markdown,
            'qti': exporter.export_to_qti,
            'moodle': exporter.export_to_moodle_xml,
            'canvas': exporter.export_to_canvas,
            'google_forms': exporter.export_to_google_forms
        }
        
        if format_type not in format_methods:
            raise ValueError(f"Unsupported format: {format_type}. Supported formats: {list(format_methods.keys())}")
        
        return format_methods[format_type](assessment, output_path)
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported export formats"""
        return ['json', 'pdf', 'word', 'text', 'markdown', 'qti', 'moodle', 'canvas', 'google_forms']
    
    def validate_output_path(self, output_path: str, format_type: str) -> str:
        """Validate and correct output path extension"""
        
        extensions = {
            'json': '.json',
            'pdf': '.pdf',
            'word': '.docx',
            'text': '.txt',
            'markdown': '.md',
            'qti': '.xml',
            'moodle': '.xml',
            'canvas': '.csv',
            'google_forms': '.csv'
        }
        
        expected_ext = extensions.get(format_type, '')
        if expected_ext and not output_path.endswith(expected_ext):
            output_path = str(Path(output_path).with_suffix(expected_ext))
        
        return output_path

# Export singleton instance
export_manager = ExportManager()
